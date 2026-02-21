"""
CogniFy Text Extraction Service
Extracts text from PDF, DOCX, TXT, XLSX, and images (OCR)
Created with love by Angela & David - 1 January 2026
"""

import os
import re
import tempfile
from typing import List, Tuple

from app.core.logging import logger
from app.domain.entities.document import FileType


class TextExtractor:
    """Extract text from various document formats"""

    @staticmethod
    def _fix_thai_ocr_spaces(text: str) -> str:
        """
        Remove spurious spaces between Thai characters.
        Tesseract Thai OCR commonly inserts spaces between every character:
            "ก า ร บ ริ ห า ร" → "การบริหาร"
            "ค ว า ม เ สี่ ย ง" → "ความเสี่ยง"
        """
        if not text:
            return text

        # Repeatedly remove single spaces between Thai characters
        # Thai Unicode range: \u0E00-\u0E7F (consonants, vowels, tone marks, digits)
        prev = None
        while prev != text:
            prev = text
            text = re.sub(r'([\u0E00-\u0E7F]) ([\u0E00-\u0E7F])', r'\1\2', text)

        return text

    @staticmethod
    def _fix_missing_spaces(text: str) -> str:
        """
        Fix missing spaces in extracted PDF text.

        Some PDFs (especially from slides or certain exporters) have text that
        looks fine visually but has no spaces when extracted.

        Examples:
            "HowNeuralNetworksWork" → "How Neural Networks Work"
            "1.Introduction" → "1. Introduction"
            "ฝึกencoderและdecoder" → "ฝึก encoder และ decoder"
        """
        if not text:
            return text

        # 1. Add space between Thai and English
        text = re.sub(r'([\u0E00-\u0E7F])([A-Za-z])', r'\1 \2', text)
        # 2. Add space between English and Thai
        text = re.sub(r'([A-Za-z])([\u0E00-\u0E7F])', r'\1 \2', text)
        # 3. camelCase fix
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        # 4. Numbers followed by letters
        text = re.sub(r'(\d)([A-Za-z])(?![A-Z0-9]|x\b|D\b)', r'\1 \2', text)
        # 5. Numbers followed by Thai
        text = re.sub(r'(\d\.?)([\u0E00-\u0E7F])', r'\1 \2', text)
        # 6. Period followed by capital
        text = re.sub(r'\.([A-Z])', r'. \1', text)
        # 7. Abbreviations
        text = re.sub(r'(e\.g\.|i\.e\.|etc\.)([A-Z])', r'\1 \2', text)
        # 8. Numbered lists
        text = re.sub(r'^(\d+\.)([A-Za-z])', r'\1 \2', text, flags=re.MULTILINE)
        # 9. Closing paren followed by letter
        text = re.sub(r'\)([A-Za-z\u0E00-\u0E7F])', r') \1', text)
        # 10. Letter followed by opening paren
        text = re.sub(r'([a-z\u0E00-\u0E7F])\(', r'\1 (', text)
        # 11. Colon followed by number
        text = re.sub(r':(\d)', r': \1', text)
        # 12. Collapse multiple spaces
        text = re.sub(r'  +', ' ', text)

        return text

    @staticmethod
    async def extract_pdf(file_path: str) -> Tuple[str, int, List[Tuple[int, str]]]:
        """
        Extract text from PDF using PyMuPDF.
        Hybrid approach: try text extraction per page, OCR only pages with no text.

        Returns:
            Tuple of (full_text, page_count, [(page_num, page_text), ...])
        """
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(file_path)
            total_pages = len(doc)

            pages: List[Tuple[int, str]] = []
            full_text_parts: List[str] = []
            ocr_pages: List[int] = []

            # First pass: try text extraction for every page
            for page_num in range(total_pages):
                page = doc[page_num]
                text = page.get_text("text").strip()

                has_text = False
                if text and len(text) >= 50:
                    real_chars = len(re.findall(r'[a-zA-Z\u0E00-\u0E7F]', text))
                    if real_chars / max(len(text), 1) >= 0.3:
                        has_text = True

                if has_text:
                    text = TextExtractor._fix_missing_spaces(text)
                    pages.append((page_num + 1, text))
                    full_text_parts.append(text)
                else:
                    ocr_pages.append(page_num)
                    pages.append((page_num + 1, ""))
                    full_text_parts.append("")

            # Second pass: OCR only pages that had no text (using Tesseract for speed)
            if ocr_pages:
                logger.info("OCR needed for {}/{} pages (Tesseract batch)", len(ocr_pages), total_pages)
                try:
                    import pytesseract
                    from PIL import Image
                    import io

                    for page_num in ocr_pages:
                        page = doc[page_num]
                        logger.debug("OCR page {}/{}", page_num + 1, total_pages)

                        mat = fitz.Matrix(2.0, 2.0)
                        pix = page.get_pixmap(matrix=mat)
                        img = Image.open(io.BytesIO(pix.tobytes("png")))

                        text = pytesseract.image_to_string(
                            img, lang="tha+eng",
                            config="--psm 1 --oem 3",
                        )
                        text = TextExtractor._fix_thai_ocr_spaces(text.strip())
                        text = TextExtractor._fix_missing_spaces(text)
                        pages[page_num] = (page_num + 1, text)
                        full_text_parts[page_num] = text

                except ImportError:
                    logger.warning("Tesseract not available, skipping OCR pages")
                except Exception as e:
                    logger.warning("Batch OCR error: {}", e)

            doc.close()

            full_text = "\n\n".join(full_text_parts)
            return full_text, len(pages), pages

        except ImportError:
            logger.warning("PyMuPDF not installed. Install with: pip install PyMuPDF")
            raise
        except Exception as e:
            logger.error("PDF extraction error: {}", e)
            raise

    @staticmethod
    async def _ocr_pdf(doc) -> Tuple[List[Tuple[int, str]], List[str]]:
        """
        OCR all pages of a PDF using Typhoon-OCR via Ollama.
        Supports Thai + English text.
        """
        import fitz  # PyMuPDF

        from app.services.ocr_service import get_ocr_service
        ocr_service = get_ocr_service()

        pages: List[Tuple[int, str]] = []
        full_text_parts: List[str] = []
        total_pages = len(doc)

        for page_num in range(total_pages):
            page = doc[page_num]
            logger.debug("OCR page {}/{}", page_num + 1, total_pages)

            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                temp_path = tmp.name
                pix.save(temp_path)

            try:
                result = await ocr_service.extract_text(temp_path)
                text = TextExtractor._fix_missing_spaces(result.text)
                pages.append((page_num + 1, text))
                full_text_parts.append(text)
            except Exception as e:
                logger.warning("OCR failed for page {}: {}", page_num + 1, e)
                pages.append((page_num + 1, ""))
                full_text_parts.append("")
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)

        return pages, full_text_parts

    @staticmethod
    async def extract_docx(file_path: str) -> Tuple[str, int, List[Tuple[int, str]]]:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            full_text = "\n\n".join(paragraphs)
            estimated_pages = max(1, len(full_text) // 3000)
            pages = [(1, full_text)]

            return full_text, estimated_pages, pages

        except ImportError:
            logger.warning("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error("DOCX extraction error: {}", e)
            raise

    @staticmethod
    async def extract_txt(file_path: str) -> Tuple[str, int, List[Tuple[int, str]]]:
        """Extract text from plain text file."""
        try:
            encodings = ['utf-8', 'utf-8-sig', 'tis-620', 'cp874', 'latin-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        full_text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError(f"Could not decode file with any encoding: {encodings}")

            estimated_pages = max(1, len(full_text) // 3000)
            pages = [(1, full_text)]
            return full_text, estimated_pages, pages

        except Exception as e:
            logger.error("TXT extraction error: {}", e)
            raise

    @staticmethod
    async def extract_xlsx(file_path: str) -> Tuple[str, int, List[Tuple[int, str]]]:
        """Extract text from Excel file."""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(file_path, read_only=True, data_only=True)
            sheets: List[Tuple[int, str]] = []
            full_text_parts: List[str] = []

            for sheet_num, sheet_name in enumerate(wb.sheetnames, 1):
                ws = wb[sheet_name]
                sheet_lines = [f"## Sheet: {sheet_name}\n"]

                for row in ws.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(v.strip() for v in row_values):
                        sheet_lines.append(" | ".join(row_values))

                sheet_text = "\n".join(sheet_lines)
                sheets.append((sheet_num, sheet_text))
                full_text_parts.append(sheet_text)

            wb.close()
            full_text = "\n\n".join(full_text_parts)
            return full_text, len(sheets), sheets

        except ImportError:
            logger.warning("openpyxl not installed. Install with: pip install openpyxl")
            raise
        except Exception as e:
            logger.error("Excel extraction error: {}", e)
            raise

    @staticmethod
    async def extract_image(file_path: str) -> Tuple[str, int, List[Tuple[int, str]]]:
        """Extract text from image using OCR (Typhoon-OCR, fallback: Tesseract)."""
        try:
            from app.services.ocr_service import get_ocr_service

            ocr_service = get_ocr_service()
            result = await ocr_service.extract_text(file_path)

            full_text = result.text
            logger.info("OCR completed: {} chars, {:.1%} confidence, engine: {}",
                        len(full_text), result.confidence, result.engine)

            pages = [(1, full_text)]
            return full_text, 1, pages

        except ImportError as e:
            logger.warning("OCR dependencies not installed: {}", e)
            raise ValueError(
                "OCR is not available. Ensure Typhoon-OCR model is pulled: ollama pull scb10x/typhoon-ocr1.5-3b"
            )
        except Exception as e:
            logger.error("Image OCR error: {}", e)
            raise

    @staticmethod
    async def extract_pdf_with_ocr(file_path: str) -> Tuple[str, int, List[Tuple[int, str]]]:
        """Extract text from scanned PDF using OCR."""
        try:
            from app.services.ocr_service import get_ocr_service

            ocr_service = get_ocr_service()
            full_text, page_results = await ocr_service.extract_from_pdf_images(file_path, dpi=300)

            pages = [(p['page'], p['text']) for p in page_results]

            avg_confidence = sum(p['confidence'] for p in page_results) / len(page_results) if page_results else 0
            logger.info("PDF OCR completed: {} pages, {:.1%} avg confidence", len(pages), avg_confidence)

            return full_text, len(pages), pages

        except ImportError as e:
            logger.warning("OCR dependencies not installed: {}", e)
            raise
        except Exception as e:
            logger.error("PDF OCR error: {}", e)
            raise

    @classmethod
    async def extract(
        cls,
        file_path: str,
        file_type: FileType,
        use_ocr_fallback: bool = True,
    ) -> Tuple[str, int, List[Tuple[int, str]]]:
        """
        Extract text from document based on file type.

        Args:
            file_path: Path to the file
            file_type: Type of file
            use_ocr_fallback: If True, use OCR for scanned PDFs with no text

        Returns:
            Tuple of (full_text, page_count, [(page_num, page_text), ...])
        """
        # Image files - always use OCR
        if file_type in (FileType.PNG, FileType.JPG, FileType.JPEG):
            logger.info("Processing image file with OCR: {}", file_path)
            return await cls.extract_image(file_path)

        # PDF files - try text extraction first, fallback to OCR
        if file_type == FileType.PDF:
            full_text, page_count, pages = await cls.extract_pdf(file_path)

            if use_ocr_fallback and not full_text.strip():
                logger.info("PDF appears to be scanned, using OCR fallback")
                return await cls.extract_pdf_with_ocr(file_path)

            return full_text, page_count, pages

        # Other document types
        if file_type in (FileType.DOCX, FileType.DOC):
            return await cls.extract_docx(file_path)
        elif file_type == FileType.TXT:
            return await cls.extract_txt(file_path)
        elif file_type in (FileType.XLSX, FileType.XLS):
            return await cls.extract_xlsx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
