"""
CogniFy Document Service
Document processing pipeline: Extract → Chunk → Embed → Store
Now with OCR support for images and scanned PDFs!
Created with love by Angela & David - 1 January 2026
"""

import os
import asyncio
import logging
from typing import Optional, List, Tuple, Dict, Any
from uuid import UUID
from datetime import datetime
from pathlib import Path

from app.core.config import settings
from app.domain.entities.document import Document, DocumentChunk, ProcessingStatus, FileType
from app.infrastructure.repositories.document_repository import DocumentRepository, DocumentChunkRepository
from app.services.embedding_service import get_embedding_service
from app.services.chunking_service import get_chunking_service, Chunk

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text from various document formats"""

    @staticmethod
    async def extract_pdf(file_path: str) -> Tuple[str, int, List[Tuple[int, str]]]:
        """
        Extract text from PDF using PyMuPDF.
        Falls back to OCR (Tesseract) for scanned/image PDFs.

        Returns:
            Tuple of (full_text, page_count, [(page_num, page_text), ...])
        """
        try:
            import fitz  # PyMuPDF
            import re

            doc = fitz.open(file_path)
            total_pages = len(doc)
            needs_ocr_count = 0

            # First pass: check if PDF needs OCR
            for page_num in range(min(5, total_pages)):  # Check first 5 pages
                page = doc[page_num]
                text = page.get_text("text").strip()

                # Check if page has meaningful text
                if not text or len(text) < 50:
                    needs_ocr_count += 1
                    continue

                # Check if text looks like real content (has Thai/English letters)
                real_chars = len(re.findall(r'[a-zA-Z\u0E00-\u0E7F]', text))
                if real_chars / max(len(text), 1) < 0.3:  # Less than 30% real chars = garbage
                    needs_ocr_count += 1

            # If more than half of checked pages need OCR, use OCR for all
            if needs_ocr_count > min(5, total_pages) / 2:
                print(f"📷 Detected scanned PDF ({total_pages} pages), using OCR...")
                pages, full_text_parts = await TextExtractor._ocr_pdf(doc)
            else:
                # Normal text extraction
                pages: List[Tuple[int, str]] = []
                full_text_parts = []
                for page_num in range(total_pages):
                    page = doc[page_num]
                    text = page.get_text("text").strip()
                    pages.append((page_num + 1, text))
                    full_text_parts.append(text)

            doc.close()

            full_text = "\n\n".join(full_text_parts)
            return full_text, len(pages), pages

        except ImportError:
            print("⚠️ PyMuPDF not installed. Install with: pip install PyMuPDF")
            raise
        except Exception as e:
            print(f"❌ PDF extraction error: {e}")
            raise

    @staticmethod
    async def _ocr_pdf(doc) -> Tuple[List[Tuple[int, str]], List[str]]:
        """
        OCR all pages of a PDF using Tesseract.
        Supports Thai + English text.
        """
        try:
            import fitz  # PyMuPDF
            import pytesseract
            from PIL import Image
            import io
            import re

            pages: List[Tuple[int, str]] = []
            full_text_parts = []
            total_pages = len(doc)

            for page_num in range(total_pages):
                page = doc[page_num]
                print(f"   🔍 OCR page {page_num + 1}/{total_pages}...")

                # Render page to image (higher resolution for better OCR)
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                pix = page.get_pixmap(matrix=mat)

                # Convert to PIL Image
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))

                # OCR with Thai + English
                text = pytesseract.image_to_string(
                    img,
                    lang='tha+eng',
                    config='--psm 1 --oem 3'  # Auto page segmentation, best OCR engine
                )

                # Clean up OCR noise from mobile screenshots
                text = TextExtractor._clean_ocr_text(text)

                pages.append((page_num + 1, text))
                full_text_parts.append(text)

            return pages, full_text_parts

        except ImportError:
            print("⚠️ pytesseract not installed. Install with: pip install pytesseract")
            print("   Also install Tesseract: brew install tesseract tesseract-lang")
            raise
        except Exception as e:
            print(f"❌ OCR error: {e}")
            raise

    @staticmethod
    def _clean_ocr_text(text: str) -> str:
        """Clean OCR noise from mobile screenshots"""
        import re

        # Remove common mobile status bar patterns
        patterns_to_remove = [
            r'\d{1,2}:\d{2}\s+\S{2,4}\s+\d{1,2}\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)',  # 13:30 ทีท 26 Dec (OCR errors in day name)
            r'\d{2}:\d{2}\s+(Mon|Tue|Wed|Thu|Fri|Sat|Sun)\s+\d+\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)',  # 08:46 Mon 22 Dec
            r'all\s+\d*G?\s*[©®]\s*\d+%\s*[๐-๙๒]*',  # all 5G © 100% ๒
            r'all\s+5G\s*[©®]?\s*\d*%?\s*[๐-๙]*',  # all 5G variations
            r'al\s+[A-Z]{1,2}\s+\d+%\s*[ส๐-๙]*',  # al FS 98% ส๒, al SF 97% ส๒
            r'^al\s+[A-Z]{1,3}\s+\d+%',  # al FS 98% at start of line
            r'\d+%\s*[ส๐-๙๒]*\s*$',  # 98% ส๒ at end
            r'^[\s\d:]+$',  # Lines with only numbers/time
            r'[©®™]+',  # Copyright symbols
        ]

        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            cleaned_line = line.strip()

            # Apply pattern removal
            for pattern in patterns_to_remove:
                cleaned_line = re.sub(pattern, '', cleaned_line, flags=re.IGNORECASE)

            # Remove excessive whitespace
            cleaned_line = re.sub(r'\s{3,}', '  ', cleaned_line)
            cleaned_line = cleaned_line.strip()

            # Only keep lines with meaningful content
            if cleaned_line and len(cleaned_line) > 2:
                # Check if line has actual letters (not just symbols)
                if re.search(r'[a-zA-Z\u0E00-\u0E7F]', cleaned_line):
                    cleaned_lines.append(cleaned_line)

        return '\n'.join(cleaned_lines)

    @staticmethod
    async def extract_docx(file_path: str) -> Tuple[str, int, List[Tuple[int, str]]]:
        """
        Extract text from DOCX using python-docx.

        Returns:
            Tuple of (full_text, page_count, [(page_num, page_text), ...])
        """
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            full_text = "\n\n".join(paragraphs)

            # DOCX doesn't have clear page breaks, estimate 1 page per 3000 chars
            estimated_pages = max(1, len(full_text) // 3000)
            pages = [(1, full_text)]  # Treat as single page

            return full_text, estimated_pages, pages

        except ImportError:
            print("⚠️ python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            print(f"❌ DOCX extraction error: {e}")
            raise

    @staticmethod
    async def extract_txt(file_path: str) -> Tuple[str, int, List[Tuple[int, str]]]:
        """
        Extract text from plain text file.

        Returns:
            Tuple of (full_text, page_count, [(page_num, page_text), ...])
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'tis-620', 'cp874', 'latin-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        full_text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError(f"Could not decode file with any encoding: {encodings}")

            # Estimate pages (1 page per 3000 chars)
            estimated_pages = max(1, len(full_text) // 3000)
            pages = [(1, full_text)]

            return full_text, estimated_pages, pages

        except Exception as e:
            print(f"❌ TXT extraction error: {e}")
            raise

    @staticmethod
    async def extract_xlsx(file_path: str) -> Tuple[str, int, List[Tuple[int, str]]]:
        """
        Extract text from Excel file.

        Returns:
            Tuple of (full_text, sheet_count, [(sheet_num, sheet_text), ...])
        """
        try:
            from openpyxl import load_workbook

            wb = load_workbook(file_path, read_only=True, data_only=True)
            sheets: List[Tuple[int, str]] = []
            full_text_parts = []

            for sheet_num, sheet_name in enumerate(wb.sheetnames, 1):
                ws = wb[sheet_name]
                sheet_lines = [f"## Sheet: {sheet_name}\n"]

                for row in ws.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(v.strip() for v in row_values):
                        sheet_lines.append(" | ".join(row_values))

                sheet_text = "\n".join(sheet_lines)
                sheets.append((sheet_num, sheet_text))
                full_text_parts.append(sheet_text)

            wb.close()

            full_text = "\n\n".join(full_text_parts)
            return full_text, len(sheets), sheets

        except ImportError:
            print("⚠️ openpyxl not installed. Install with: pip install openpyxl")
            raise
        except Exception as e:
            print(f"❌ Excel extraction error: {e}")
            raise

    @staticmethod
    async def extract_image(file_path: str) -> Tuple[str, int, List[Tuple[int, str]]]:
        """
        Extract text from image using OCR.

        Supports: PNG, JPG, JPEG
        Uses multiple OCR engines with fallback (Tesseract → PaddleOCR → EasyOCR)

        Returns:
            Tuple of (full_text, page_count, [(page_num, page_text), ...])
        """
        try:
            from app.services.ocr_service import get_ocr_service

            ocr_service = get_ocr_service()
            result = await ocr_service.extract_text(file_path, preprocess=True)

            full_text = result.text
            confidence = result.confidence

            logger.info(f"🔍 OCR completed: {len(full_text)} chars, {confidence:.1%} confidence, engine: {result.engine}")

            # Return as single page
            pages = [(1, full_text)]
            return full_text, 1, pages

        except ImportError as e:
            logger.warning(f"⚠️ OCR dependencies not installed: {e}")
            raise ValueError(
                "OCR is not available. Install with: pip install pytesseract pillow opencv-python\n"
                "Also install Tesseract: brew install tesseract tesseract-lang"
            )
        except Exception as e:
            logger.error(f"❌ Image OCR error: {e}")
            raise

    @staticmethod
    async def extract_pdf_with_ocr(file_path: str) -> Tuple[str, int, List[Tuple[int, str]]]:
        """
        Extract text from scanned PDF using OCR.

        This is used when regular PDF extraction returns no text (scanned document).

        Returns:
            Tuple of (full_text, page_count, [(page_num, page_text), ...])
        """
        try:
            from app.services.ocr_service import get_ocr_service

            ocr_service = get_ocr_service()
            full_text, page_results = await ocr_service.extract_from_pdf_images(file_path, dpi=300)

            pages = [(p['page'], p['text']) for p in page_results]

            avg_confidence = sum(p['confidence'] for p in page_results) / len(page_results) if page_results else 0
            logger.info(f"🔍 PDF OCR completed: {len(pages)} pages, {avg_confidence:.1%} avg confidence")

            return full_text, len(pages), pages

        except ImportError as e:
            logger.warning(f"⚠️ OCR dependencies not installed: {e}")
            raise
        except Exception as e:
            logger.error(f"❌ PDF OCR error: {e}")
            raise

    @classmethod
    async def extract(
        cls,
        file_path: str,
        file_type: FileType,
        use_ocr_fallback: bool = True,
    ) -> Tuple[str, int, List[Tuple[int, str]]]:
        """
        Extract text from document based on file type.

        Args:
            file_path: Path to the file
            file_type: Type of file
            use_ocr_fallback: If True, use OCR for scanned PDFs with no text

        Returns:
            Tuple of (full_text, page_count, [(page_num, page_text), ...])
        """
        # Image files - always use OCR
        if file_type in (FileType.PNG, FileType.JPG, FileType.JPEG):
            logger.info(f"🖼️ Processing image file with OCR: {file_path}")
            return await cls.extract_image(file_path)

        # PDF files - try text extraction first, fallback to OCR
        if file_type == FileType.PDF:
            full_text, page_count, pages = await cls.extract_pdf(file_path)

            # Check if PDF is scanned (no text)
            if use_ocr_fallback and not full_text.strip():
                logger.info(f"📄 PDF appears to be scanned, using OCR fallback...")
                return await cls.extract_pdf_with_ocr(file_path)

            return full_text, page_count, pages

        # Other document types
        if file_type in (FileType.DOCX, FileType.DOC):
            return await cls.extract_docx(file_path)
        elif file_type == FileType.TXT:
            return await cls.extract_txt(file_path)
        elif file_type in (FileType.XLSX, FileType.XLS):
            return await cls.extract_xlsx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")


class DocumentService:
    """
    Document processing service.

    Pipeline:
    1. Extract text from document
    2. Chunk text with overlap
    3. Generate embeddings for chunks
    4. Store chunks in database
    """

    def __init__(self):
        self.document_repo = DocumentRepository()
        self.chunk_repo = DocumentChunkRepository()
        self.embedding_service = get_embedding_service()
        self.chunking_service = get_chunking_service()

    async def process_document(
        self,
        document_id: UUID,
        on_progress: Optional[callable] = None
    ) -> Document:
        """
        Process a document: extract → chunk → embed → store.

        Args:
            document_id: ID of the document to process
            on_progress: Optional callback for progress updates

        Returns:
            Updated Document entity
        """
        # Get document
        document = await self.document_repo.get_by_id(document_id)
        if document is None:
            raise ValueError(f"Document not found: {document_id}")

        if not document.file_path or not os.path.exists(document.file_path):
            await self._fail_document(document_id, "File not found")
            raise FileNotFoundError(f"File not found: {document.file_path}")

        try:
            # Update status to processing
            await self.document_repo.update_status(document_id, ProcessingStatus.PROCESSING)

            if on_progress:
                await on_progress("extracting", 0)

            # 1. Extract text
            print(f"📄 Extracting text from {document.original_filename}...")
            full_text, page_count, pages = await TextExtractor.extract(
                document.file_path,
                document.file_type
            )

            if not full_text.strip():
                await self._fail_document(document_id, "No text content found")
                raise ValueError("No text content found in document")

            # Update page count
            document.page_count = page_count

            if on_progress:
                await on_progress("chunking", 20)

            # 2. Chunk text
            print(f"✂️ Chunking text into segments...")
            if len(pages) > 1:
                chunks = self.chunking_service.chunk_by_pages(pages)
            else:
                chunks = self.chunking_service.chunk_text(full_text)

            print(f"   Created {len(chunks)} chunks")

            if on_progress:
                await on_progress("embedding", 40)

            # 3. Generate embeddings
            print(f"🧠 Generating embeddings...")
            chunk_texts = [chunk.content for chunk in chunks]
            embeddings = await self.embedding_service.get_embeddings_batch(
                chunk_texts,
                batch_size=5
            )

            # Count successful embeddings
            successful = sum(1 for e in embeddings if e is not None)
            print(f"   Generated {successful}/{len(chunks)} embeddings")

            if on_progress:
                await on_progress("storing", 80)

            # 4. Create DocumentChunk entities
            print(f"💾 Storing chunks...")
            document_chunks: List[DocumentChunk] = []

            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                doc_chunk = DocumentChunk(
                    document_id=document_id,
                    chunk_index=i,
                    content=chunk.content,
                    page_number=chunk.page_number,
                    section_title=chunk.section_title,
                    token_count=chunk.token_count,
                    embedding=embedding,
                )
                document_chunks.append(doc_chunk)

            # 5. Delete existing chunks and insert new ones
            await self.chunk_repo.delete_by_document(document_id)
            await self.chunk_repo.create_batch(document_chunks)

            # 6. Update document status to COMPLETED
            document.processing_status = ProcessingStatus.COMPLETED
            document.page_count = page_count
            document.total_chunks = len(document_chunks)
            document.processed_at = datetime.now()
            await self.document_repo.update(document)

            if on_progress:
                await on_progress("completed", 100)

            print(f"✅ Document processed successfully: {document.original_filename}")
            print(f"   Pages: {page_count}, Chunks: {len(document_chunks)}")

            return await self.document_repo.get_by_id(document_id)

        except Exception as e:
            print(f"❌ Document processing error: {e}")
            await self._fail_document(document_id, str(e))
            raise

    async def _fail_document(self, document_id: UUID, error: str) -> None:
        """Mark document as failed"""
        await self.document_repo.update_status(
            document_id,
            ProcessingStatus.FAILED,
            error=error
        )

    async def reprocess_document(self, document_id: UUID) -> Document:
        """Reprocess an existing document"""
        # Delete existing chunks
        await self.chunk_repo.delete_by_document(document_id)

        # Reset status
        await self.document_repo.update_status(document_id, ProcessingStatus.PENDING)

        # Process again
        return await self.process_document(document_id)

    async def get_document_stats(self, document_id: UUID) -> Dict[str, Any]:
        """Get statistics for a document"""
        document = await self.document_repo.get_by_id(document_id)
        if document is None:
            raise ValueError(f"Document not found: {document_id}")

        chunks = await self.chunk_repo.get_by_document(document_id)

        total_tokens = sum(c.token_count or 0 for c in chunks)
        chunks_with_embeddings = sum(1 for c in chunks if c.embedding)

        return {
            "document_id": str(document_id),
            "filename": document.original_filename,
            "status": document.processing_status.value,
            "page_count": document.page_count,
            "total_chunks": len(chunks),
            "chunks_with_embeddings": chunks_with_embeddings,
            "total_tokens": total_tokens,
            "avg_tokens_per_chunk": total_tokens // len(chunks) if chunks else 0,
        }


# ============================================================================
# BACKGROUND PROCESSING
# ============================================================================

def process_document_background(document_id: UUID) -> None:
    """
    Background task to process a document.
    Uses a completely separate database connection to avoid event loop issues.
    """
    import asyncio
    import asyncpg

    print(f"🚀 Starting background processing for document: {document_id}")

    async def _process_with_own_connection():
        """Process document with its own database connection"""
        from app.core.config import settings
        from app.domain.entities.document import ProcessingStatus
        from datetime import datetime

        # Create dedicated connection for this background task
        conn = await asyncpg.connect(settings.DATABASE_URL)

        try:
            # 1. Get document
            doc_row = await conn.fetchrow(
                "SELECT * FROM documents WHERE document_id = $1",
                document_id
            )

            if not doc_row:
                print(f"❌ Document not found: {document_id}")
                return

            file_path = doc_row['file_path']
            file_type_str = doc_row['file_type']
            original_filename = doc_row['original_filename']

            if not file_path or not os.path.exists(file_path):
                await conn.execute(
                    "UPDATE documents SET processing_status = $1, processing_error = $2 WHERE document_id = $3",
                    'failed', 'File not found', document_id
                )
                print(f"❌ File not found: {file_path}")
                return

            # 2. Update status to processing
            await conn.execute(
                "UPDATE documents SET processing_status = $1, processing_step = $2, processing_progress = $3 WHERE document_id = $4",
                'processing', 'extracting', 0, document_id
            )

            # 3. Extract text
            print(f"📄 Extracting text from {original_filename}...")
            from app.domain.entities.document import FileType
            file_type = FileType(file_type_str)
            full_text, page_count, pages = await TextExtractor.extract(file_path, file_type)

            if not full_text.strip():
                await conn.execute(
                    "UPDATE documents SET processing_status = $1, processing_error = $2 WHERE document_id = $3",
                    'failed', 'No text content found', document_id
                )
                print(f"❌ No text content found")
                return

            # 4. Chunk text
            await conn.execute(
                "UPDATE documents SET processing_step = $1, processing_progress = $2 WHERE document_id = $3",
                'chunking', 25, document_id
            )
            print(f"✂️ Chunking text into segments...")
            chunking_service = get_chunking_service()

            # Calculate average chars per page to detect slides/short content
            avg_chars_per_page = len(full_text) / max(page_count, 1)

            if avg_chars_per_page < 500 and len(pages) > 1:
                # Short content per page (likely slides) - use 1 page = 1 chunk
                print(f"   Detected slides/short content ({avg_chars_per_page:.0f} chars/page avg)")
                chunks = []
                for page_num, page_text in pages:
                    if page_text.strip():  # Skip empty pages
                        from app.services.chunking_service import Chunk
                        chunks.append(Chunk(
                            content=page_text.strip(),
                            index=len(chunks),
                            start_char=0,
                            end_char=len(page_text),
                            token_count=len(page_text.split()),
                            page_number=page_num,
                            section_title=None
                        ))
            elif len(pages) > 1:
                chunks = chunking_service.chunk_by_pages(pages)
            else:
                chunks = chunking_service.chunk_text(full_text)
            print(f"   Created {len(chunks)} chunks")

            # 5. Generate embeddings (direct Ollama call to avoid event loop issues)
            await conn.execute(
                "UPDATE documents SET processing_step = $1, processing_progress = $2 WHERE document_id = $3",
                'embedding', 40, document_id
            )
            print(f"🧠 Generating embeddings...")
            import httpx
            embeddings = []
            total_chunks = len(chunks)
            for i, chunk in enumerate(chunks):
                # Update progress for each embedding (40-90% range)
                embed_progress = 40 + int((i / total_chunks) * 50)
                await conn.execute(
                    "UPDATE documents SET processing_progress = $1 WHERE document_id = $2",
                    embed_progress, document_id
                )
                try:
                    async with httpx.AsyncClient(timeout=60.0) as http_client:
                        response = await http_client.post(
                            "http://localhost:11434/api/embeddings",
                            json={"model": "nomic-embed-text", "prompt": chunk.content[:8000]}
                        )
                        if response.status_code == 200:
                            data = response.json()
                            embeddings.append(data.get("embedding"))
                        else:
                            embeddings.append(None)
                except Exception as emb_err:
                    print(f"⚠️ Embedding error for chunk {i}: {emb_err}")
                    embeddings.append(None)
            successful = sum(1 for e in embeddings if e is not None)
            print(f"   Generated {successful}/{len(chunks)} embeddings")

            # 6. Delete existing chunks
            await conn.execute(
                "DELETE FROM document_chunks WHERE document_id = $1",
                document_id
            )

            # 7. Insert new chunks
            await conn.execute(
                "UPDATE documents SET processing_step = $1, processing_progress = $2 WHERE document_id = $3",
                'storing', 90, document_id
            )
            print(f"💾 Storing chunks...")
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                import uuid
                chunk_id = uuid.uuid4()
                # Convert embedding list to pgvector format string
                embedding_str = None
                if embedding:
                    embedding_str = '[' + ','.join(str(x) for x in embedding) + ']'
                await conn.execute(
                    """
                    INSERT INTO document_chunks
                    (chunk_id, document_id, chunk_index, content, page_number, section_title, token_count, embedding)
                    VALUES ($1, $2, $3, $4, $5, $6, $7, $8::vector)
                    """,
                    chunk_id, document_id, i, chunk.content, chunk.page_number,
                    chunk.section_title, chunk.token_count, embedding_str
                )

            # 8. Update document as completed
            await conn.execute(
                """
                UPDATE documents
                SET processing_status = $1, processing_step = $2, processing_progress = $3,
                    page_count = $4, total_chunks = $5, processed_at = $6
                WHERE document_id = $7
                """,
                'completed', 'completed', 100, page_count, len(chunks), datetime.now(), document_id
            )

            print(f"✅ Document processed successfully: {original_filename}")
            print(f"   Pages: {page_count}, Chunks: {len(chunks)}")

        except Exception as e:
            print(f"❌ Processing error: {e}")
            import traceback
            traceback.print_exc()
            await conn.execute(
                "UPDATE documents SET processing_status = $1, processing_error = $2 WHERE document_id = $3",
                'failed', str(e)[:500], document_id
            )
        finally:
            await conn.close()

    # Run in new event loop
    try:
        asyncio.run(_process_with_own_connection())
    except Exception as e:
        print(f"❌ Background task failed: {e}")
        import traceback
        traceback.print_exc()


# ============================================================================
# SINGLETON INSTANCE
# ============================================================================

_document_service: Optional[DocumentService] = None


def get_document_service() -> DocumentService:
    """Get global DocumentService instance"""
    global _document_service
    if _document_service is None:
        _document_service = DocumentService()
    return _document_service
